"""
Document Parser
Parses various document formats (Word, PDF) and extracts text content.
"""

import os
import logging
from typing import Dict, List, Any, Optional

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for various document formats."""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc', '.pdf']
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document and extract text content.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Dict containing extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Get file extension
        _, ext = os.path.splitext(file_path.lower())
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")
        
        logger.info(f"Parsing document: {file_path}")
        
        try:
            if ext == '.pdf':
                return self._parse_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._parse_word(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF document."""
        try:
            # Try PyPDF2 first
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                paragraphs = []
                tables = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            # Split into paragraphs
                            page_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                            paragraphs.extend(page_paragraphs)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        continue
                
                # Count words
                total_words = sum(len(p.split()) for p in paragraphs)
                
                logger.info(f"Successfully parsed PDF: {file_path}")
                logger.info(f"Extracted {len(paragraphs)} paragraphs, {total_words} words, {len(tables)} tables")
                
                return {
                    'paragraphs': paragraphs,
                    'tables': tables,
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'pdf',
                        'num_pages': len(pdf_reader.pages),
                        'num_paragraphs': len(paragraphs),
                        'num_words': total_words,
                        'num_tables': len(tables)
                    }
                }
                
        except ImportError:
            # Fallback to pdfplumber
            try:
                import pdfplumber
                
                paragraphs = []
                tables = []
                
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            # Extract text
                            text = page.extract_text()
                            if text:
                                page_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                                paragraphs.extend(page_paragraphs)
                            
                            # Extract tables
                            page_tables = page.extract_tables()
                            if page_tables:
                                for table in page_tables:
                                    if table:
                                        tables.append({
                                            'page': page_num + 1,
                                            'data': table
                                        })
                        except Exception as e:
                            logger.warning(f"Error extracting from page {page_num}: {e}")
                            continue
                
                total_words = sum(len(p.split()) for p in paragraphs)
                
                logger.info(f"Successfully parsed PDF with pdfplumber: {file_path}")
                logger.info(f"Extracted {len(paragraphs)} paragraphs, {total_words} words, {len(tables)} tables")
                
                return {
                    'paragraphs': paragraphs,
                    'tables': tables,
                    'metadata': {
                        'file_path': file_path,
                        'file_type': 'pdf',
                        'num_pages': len(pdf.pages),
                        'num_paragraphs': len(paragraphs),
                        'num_words': total_words,
                        'num_tables': len(tables)
                    }
                }
                
            except ImportError:
                # Final fallback to pymupdf
                try:
                    import fitz  # PyMuPDF
                    
                    doc = fitz.open(file_path)
                    paragraphs = []
                    tables = []
                    
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text = page.get_text()
                        
                        if text.strip():
                            page_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                            paragraphs.extend(page_paragraphs)
                    
                    doc.close()
                    
                    total_words = sum(len(p.split()) for p in paragraphs)
                    
                    logger.info(f"Successfully parsed PDF with PyMuPDF: {file_path}")
                    logger.info(f"Extracted {len(paragraphs)} paragraphs, {total_words} words, {len(tables)} tables")
                    
                    return {
                        'paragraphs': paragraphs,
                        'tables': tables,
                        'metadata': {
                            'file_path': file_path,
                            'file_type': 'pdf',
                            'num_pages': len(doc),
                            'num_paragraphs': len(paragraphs),
                            'num_words': total_words,
                            'num_tables': len(tables)
                        }
                    }
                    
                except ImportError:
                    raise ImportError("No PDF parsing library available. Please install PyPDF2, pdfplumber, or PyMuPDF")
    
    def _parse_word(self, file_path: str) -> Dict[str, Any]:
        """Parse Word document (.docx or .doc)."""
        try:
            from docx import Document
            
            # Load document
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:  # Only add non-empty tables
                    tables.append({
                        'index': table_idx,
                        'data': table_data
                    })
            
            # Count words
            total_words = sum(len(p.split()) for p in paragraphs)
            
            logger.info(f"Successfully parsed document: {file_path}")
            logger.info(f"Extracted {len(paragraphs)} paragraphs, {total_words} words, {len(tables)} tables")
            
            return {
                'paragraphs': paragraphs,
                'tables': tables,
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'docx',
                    'num_paragraphs': len(paragraphs),
                    'num_words': total_words,
                    'num_tables': len(tables)
                }
            }
            
        except ImportError:
            raise ImportError("python-docx library not available. Please install it to parse Word documents.")
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return self.supported_formats.copy()


def main():
    """Test the document parser."""
    parser = DocumentParser()
    
    # Test with a sample file (if available)
    test_files = [
        "/home/ubuntu/llm-evaluation/data/raw/source.pdf",
        "/home/ubuntu/llm-evaluation/data/raw/source.docx"
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            try:
                result = parser.parse_document(test_file)
                print(f"\nParsed {test_file}:")
                print(f"- Paragraphs: {len(result['paragraphs'])}")
                print(f"- Tables: {len(result['tables'])}")
                print(f"- Words: {result['metadata']['num_words']}")
                
                # Show first few paragraphs
                if result['paragraphs']:
                    print(f"- First paragraph: {result['paragraphs'][0][:100]}...")
                    
            except Exception as e:
                print(f"Error parsing {test_file}: {e}")
        else:
            print(f"Test file not found: {test_file}")


if __name__ == "__main__":
    main()
